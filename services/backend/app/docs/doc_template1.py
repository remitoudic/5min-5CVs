import random
from docx import Document
from docx.shared import Pt

class TemplateCv1:
    def __init__(self, filename):
        self.filename = filename
        self.doc = Document()

    def add_name(self, name):
        title = self.doc.add_paragraph()
        run = title.add_run(name)
        run.bold = True
        run.font.size = Pt(24)
        self.doc.add_paragraph()  # Add spacing

    def add_career_objectives(self, objectives):
        self.doc.add_heading("Career Objectives", level=1)
        self.doc.add_paragraph(objectives)

    def add_work_experience(self, work_experience):
        self.doc.add_heading("Work Experience", level=1)
        table = self.doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Position'
        hdr_cells[1].text = 'Company'
        hdr_cells[2].text = 'Date'
        for exp in work_experience:
            row_cells = table.add_row().cells
            row_cells[0].text = exp["position"]
            row_cells[1].text = exp["company"]
            row_cells[2].text = exp["date"]

    def add_education(self, education):
        self.doc.add_heading("Education", level=1)
        table = self.doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'School'
        hdr_cells[1].text = 'Date'
        for edu in education:
            row_cells = table.add_row().cells
            row_cells[0].text = edu["school"]
            row_cells[1].text = edu["date"]

    def build_doc(self, data):
        self.add_name(data.get("name", ""))
        self.add_career_objectives(data.get("career_objectives", ""))
        self.add_work_experience(data.get("work_experience", []))
        self.add_education(data.get("education", []))
        self.doc.save(self.filename)

def generate_random_id():
    return random.randint(1000, 9999)

# Example usage
data = {
    "name": "Joe Doe",
    "career_objectives": "Seeking a challenging role in data science where I can utilize my skills to contribute to organizational success.",
    "work_experience": [
        {"position": "Data Analyst", "company": "DataCorp", "date": "2019 - Present"},
        {"position": "Junior Data Scientist", "company": "Tech Solutions", "date": "2017 - 2019"}
    ],
    "education": [
        {"school": "University of Data Science", "date": "2013 - 2017"},
        {"school": "San Joseph High School", "date": "2010 - 2012"}
    ]
}

# Generate a random ID
random_id = generate_random_id()

# Create the filename using the template "template1_" + name + "_" + random_id
filename = f"template1_{data['name'].replace(' ', '_')}_{random_id}.docx"

# Create the CV Word document using the class
cv = TemplateCv1(filename)
cv.build_doc(data)

print(f"CV saved as {filename}")
